# -*- coding: utf-8 -*-
"""Genera el documento Word 'Como se descuenta el tiempo' con graficas."""
import os
from datetime import datetime, timedelta

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.dirname(os.path.abspath(__file__))
IMG = os.path.join(OUT, "_img")
os.makedirs(IMG, exist_ok=True)

# Colores
C_TRABAJO = "#2E7D32"   # verde
C_PAUSA   = "#C62828"   # rojo
C_PLANEADO= "#1565C0"   # azul
C_EXTRA   = "#F9A825"   # amarillo
C_FALTA   = "#6A1B9A"   # morado
C_BOX     = "#ECEFF1"
C_LINE    = "#37474F"

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 11,
    "axes.edgecolor": C_LINE,
})


def to_min(h, m=0):
    return h * 60 + m


# ─────────────────────────────────────────────────────────────────────────────
# Diagrama 1: linea de tiempo - caso Marcaje de Reloj (11:10 / 18:43, sin par)
# ─────────────────────────────────────────────────────────────────────────────
def d1_timeline_marcaje():
    fig, ax = plt.subplots(figsize=(11, 3.2))
    start = to_min(11, 10)
    end = to_min(18, 43)
    # barra de trabajo continuo
    ax.barh(0, end - start, left=start, height=0.5, color=C_TRABAJO, edgecolor="none")
    # marcas
    for x, lbl, off in [(start, "Entrada\n11:10", -25), (end, "Salida\n18:43", 25)]:
        ax.plot(x, 0, "o", color=C_LINE, markersize=9)
        ax.annotate(lbl, (x, 0), xytext=(x + off, 0.55), ha="center", va="bottom",
                    fontsize=10, fontweight="bold",
                    arrowprops=dict(arrowstyle="-", color=C_LINE, lw=0.8))
    # ventana del descanso planeado D1 (14:00-14:15) NO tomado
    d1i, d1f = to_min(14, 0), to_min(14, 15)
    ax.axvspan(d1i, d1f, ymin=0.05, ymax=0.45, color=C_PAUSA, alpha=0.18, hatch="//")
    ax.annotate("Descanso planeado D1\n14:00-14:15 (NO marcado)\n→ no se descuenta\n(trabajo continuo)",
                (d1i + 7, 0), xytext=(d1i + 7, -0.85), ha="center", va="top",
                fontsize=9, color=C_PAUSA,
                arrowprops=dict(arrowstyle="->", color=C_PAUSA, lw=1))
    # etiqueta bruto
    ax.annotate("Bruto = 453 min\n(11:10 → 18:43)", ((start + end) / 2, 0),
                xytext=((start + end) / 2, -0.25), ha="center", va="top",
                fontsize=10, fontweight="bold", color="white")
    ax.set_xlim(start - 30, end + 30)
    ax.set_ylim(-1.5, 1.1)
    ax.set_yticks([])
    # eje horario
    ticks = [to_min(11), to_min(12), to_min(13), to_min(14), to_min(15),
             to_min(16), to_min(17), to_min(18), to_min(19)]
    ax.set_xticks(ticks)
    ax.set_xticklabels([f"{h}:00" for h in range(11, 20)])
    ax.set_title("Marcaje de Reloj — caso 11:10 / 18:43 (sin par intermedio)",
                 fontsize=12, fontweight="bold", pad=10)
    for s in ["top", "right", "left"]:
        ax.spines[s].set_visible(False)
    plt.tight_layout()
    p = os.path.join(IMG, "d1_timeline_marcaje.png")
    plt.savefig(p, dpi=150, bbox_inches="tight")
    plt.close()
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Diagrama 2: linea de tiempo - caso Aralim (con pausas reales)
# ─────────────────────────────────────────────────────────────────────────────
def d2_timeline_aralim():
    fig, ax = plt.subplots(figsize=(11, 3.6))
    start = to_min(7, 9)
    end = to_min(18, 18)
    # segmentos trabajados
    segs = [(to_min(7, 9), to_min(11, 3)),
            (to_min(11, 30), to_min(14, 0)),
            (to_min(14, 44), to_min(18, 18))]
    for s, e in segs:
        ax.barh(0, e - s, left=s, height=0.5, color=C_TRABAJO, edgecolor="none")
    # pausas
    pauses = [(to_min(11, 3), to_min(11, 30), "D1 marcado\n27 min"),
              (to_min(14, 0), to_min(14, 44), "D2 par intermedio\n44 min (pausa por defecto)")]
    for s, e, lbl in pauses:
        ax.barh(0, e - s, left=s, height=0.5, color=C_PAUSA, alpha=0.85, edgecolor="none")
        ax.annotate(lbl, ((s + e) / 2, 0), xytext=((s + e) / 2, -0.85), ha="center",
                    va="top", fontsize=8.5, color=C_PAUSA, fontweight="bold")
    # marcas
    marks = [(to_min(7, 9), "7:09"), (to_min(11, 3), "11:03"), (to_min(11, 30), "11:30"),
             (to_min(14, 0), "14:00"), (to_min(14, 44), "14:44"), (to_min(18, 18), "18:18")]
    for x, lbl in marks:
        ax.plot(x, 0, "o", color=C_LINE, markersize=6)
        ax.text(x, 0.32, lbl, ha="center", va="bottom", fontsize=8.5, rotation=0)
    ax.annotate("Bruto = 669 min   |   Pausas = 27 + 44 = 71   |   Trabajado = 598 min",
                ((start + end) / 2, 0), xytext=((start + end) / 2, -1.45), ha="center",
                fontsize=10, fontweight="bold", color=C_LINE)
    ax.set_xlim(start - 20, end + 20)
    ax.set_ylim(-1.9, 1.0)
    ax.set_yticks([])
    ticks = [to_min(7), to_min(9), to_min(11), to_min(13), to_min(15), to_min(17), to_min(18, 18)]
    ax.set_xticks(ticks)
    ax.set_xticklabels(["7:00", "9:00", "11:00", "13:00", "15:00", "17:00", "18:18"])
    ax.set_title("Marcaje de Reloj — caso Aralim (con pares intermedios = pausas)",
                 fontsize=12, fontweight="bold", pad=10)
    for s in ["top", "right", "left"]:
        ax.spines[s].set_visible(False)
    plt.tight_layout()
    p = os.path.join(IMG, "d2_timeline_aralim.png")
    plt.savefig(p, dpi=150, bbox_inches="tight")
    plt.close()
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Diagrama 3: barras Trabajado vs Planeado → Acreditado + Extra
# ─────────────────────────────────────────────────────────────────────────────
def d3_bars():
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.2))

    # Caso A: 11:10/18:43 → extra
    ax = axes[0]
    trabajado = 453
    planeado = 435
    acreditado = min(trabajado, planeado)
    extra = max(0, trabajado - planeado)
    ax.bar("Tiempo\nAcreditado", acreditado, color=C_TRABAJO, edgecolor="white")
    ax.bar("Tiempo\nAcreditado", extra, bottom=acreditado, color=C_EXTRA, edgecolor="white")
    ax.axhline(planeado, color=C_PLANEADO, lw=2, ls="--")
    ax.text(0, planeado + 8, f"Planeado = {planeado}", ha="center", color=C_PLANEADO, fontweight="bold")
    ax.text(0, acreditado / 2, f"{acreditado}", ha="center", va="center", color="white", fontweight="bold")
    ax.text(0, acreditado + extra / 2, f"Extra\n{extra}", ha="center", va="center", color="black", fontweight="bold")
    ax.set_title(f"Caso 11:10/18:43\nTrabajado = {trabajado} → Acreditado {acreditado} + Extra {extra}",
                 fontsize=10, fontweight="bold")
    ax.set_ylim(0, 520)
    ax.set_ylabel("Minutos")

    # Caso B: faltante con permiso
    ax = axes[1]
    trabajado = 500
    planeado = 540
    permiso = 30
    acreditado = min(trabajado + permiso, planeado)  # 530
    faltante = max(0, planeado - trabajado - permiso)  # 10
    ax.bar("Tiempo\nAcreditado", acreditado, color=C_TRABAJO, edgecolor="white")
    # faltante restante
    if faltante:
        ax.bar("Tiempo\nAcreditado", faltante, bottom=acreditado, color=C_FALTA, alpha=0.6, edgecolor="white")
        ax.text(0, acreditado + faltante / 2, f"Faltante\n{faltante}", ha="center", va="center", color="white", fontweight="bold")
    ax.axhline(planeado, color=C_PLANEADO, lw=2, ls="--")
    ax.text(0, planeado + 8, f"Planeado = {planeado}", ha="center", color=C_PLANEADO, fontweight="bold")
    ax.text(0, acreditado / 2, f"{acreditado}", ha="center", va="center", color="white", fontweight="bold")
    ax.set_title(f"Caso con permiso\nTrabajado {trabajado} + Permiso {permiso} → Acreditado {acreditado}, Faltante {faltante}",
                 fontsize=10, fontweight="bold")
    ax.set_ylim(0, 600)
    ax.set_ylabel("Minutos")

    plt.tight_layout()
    p = os.path.join(IMG, "d3_bars.png")
    plt.savefig(p, dpi=150, bbox_inches="tight")
    plt.close()
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Diagrama 4: flujo del descuento (cajas y flechas)
# ─────────────────────────────────────────────────────────────────────────────
def d4_flujo():
    fig, ax = plt.subplots(figsize=(11, 5.2))
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    def box(x, y, w, h, text, color=C_BOX, txtcolor="black", bold=True):
        b = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.5,rounding_size=2",
                           linewidth=1.2, edgecolor=C_LINE, facecolor=color)
        ax.add_patch(b)
        ax.text(x + w / 2, y + h / 2, text, ha="center", va="center",
                fontsize=9.5, fontweight="bold" if bold else "normal",
                color=txtcolor, wrap=True)

    def arrow(x1, y1, x2, y2, label="", color=C_LINE):
        a = FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=14,
                            lw=1.4, color=color)
        ax.add_patch(a)
        if label:
            ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 2.5, label, ha="center", fontsize=8.5,
                    color=color, fontweight="bold")

    # Fila 1
    box(5, 78, 22, 14, "Bruto\n(primera → última marca)", "#BBDEFB")
    box(39, 78, 24, 14, "− Pausas no pagadas\n(par intermedio real)", "#FFCDD2")
    box(75, 78, 22, 14, "Trabajado\n(reloj)", "#C8E6C9")
    arrow(27, 85, 39, 85)
    arrow(63, 85, 75, 85)

    # Fila 2: comparacion con planeado
    box(5, 48, 22, 14, "Planeado\n(JornadaNeta o\noverride del modal)", "#BBDEFB")
    arrow(16, 78, 16, 62, "compara")
    arrow(16, 62, 86, 62)
    arrow(86, 62, 86, 78)
    box(39, 48, 24, 14, "Trabajado vs Planeado", "#FFF9C4")
    arrow(16, 55, 39, 55)

    # Fila 3: ramas extra / faltante
    box(5, 22, 24, 14, "Extra detectado\n= Max(0, Trab − Planeado)", "#FFE0B2")
    box(40, 22, 24, 14, "Faltante\n= Max(0, Planeado − Trab − Permiso)", "#E1BEE7")
    box(75, 22, 22, 14, "Tiempo Acreditado\n= Min(Trab + Permiso, Planeado)", "#C8E6C9")
    arrow(28, 48, 17, 36, "si Trab > Planeado")
    arrow(50, 48, 52, 36, "si Trab < Planeado")
    arrow(64, 29, 75, 29, "+ Permiso cubre")

    # nota umbral / neteo
    ax.text(50, 8, "Extra detectado por día  →  aprobación SEMANAL (neto de la semana)",
            ha="center", fontsize=9.5, color=C_FALTA, fontweight="bold",
            bbox=dict(boxstyle="round,pad=0.4", facecolor="#F3E5F5", edgecolor=C_FALTA))

    ax.set_title("Flujo del descuento del tiempo (ambos modos)", fontsize=13, fontweight="bold", pad=8)
    plt.tight_layout()
    p = os.path.join(IMG, "d4_flujo.png")
    plt.savefig(p, dpi=150, bbox_inches="tight")
    plt.close()
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Diagrama 5: neteo semanal
# ─────────────────────────────────────────────────────────────────────────────
def d5_neteo():
    fig, ax = plt.subplots(figsize=(11, 4.4))
    days = ["Lun", "Mar", "Mié", "Jue", "Vie"]
    trabajado = [600, 480, 540, 540, 600]
    planeado = 540
    x = np.arange(len(days))
    # barras: parte acreditada (hasta planeado) + extra o faltante
    for i, t in enumerate(trabajado):
        if t >= planeado:
            ax.bar(i, planeado, color=C_TRABAJO, edgecolor="white", width=0.6)
            ax.bar(i, t - planeado, bottom=planeado, color=C_EXTRA, edgecolor="white", width=0.6)
            ax.text(i, t + 8, f"+{t - planeado}", ha="center", color=C_EXTRA, fontweight="bold", fontsize=9)
        else:
            ax.bar(i, t, color=C_TRABAJO, edgecolor="white", width=0.6)
            ax.bar(i, planeado - t, bottom=t, color=C_FALTA, alpha=0.55, edgecolor="white", width=0.6, hatch="//")
            ax.text(i, planeado + 8, f"-{planeado - t}", ha="center", color=C_FALTA, fontweight="bold", fontsize=9)
        ax.text(i, min(t, planeado) / 2, str(t), ha="center", va="center", color="white", fontweight="bold")
    ax.axhline(planeado, color=C_PLANEADO, lw=2, ls="--")
    ax.text(len(days) - 0.5, planeado + 22, f"Planeado diario = {planeado}", ha="right", color=C_PLANEADO, fontweight="bold")
    ax.set_xticks(x)
    ax.set_xticklabels(days)
    ax.set_ylabel("Minutos trabajados")
    ax.set_ylim(0, 700)
    sum_t = sum(trabajado)
    sum_p = planeado * len(days)
    neto = max(0, sum_t - sum_p)
    ax.set_title(f"Neteo semanal — ∑Trabajado = {sum_t}, ∑Planeado = {sum_p}  →  Extra neto = {neto}  "
                 f"(el faltante del Mar compensa parte del extra)",
                 fontsize=10.5, fontweight="bold")
    # leyenda
    from matplotlib.patches import Patch
    ax.legend(handles=[Patch(color=C_TRABAJO, label="Acreditado"),
                       Patch(color=C_EXTRA, label="Extra del día"),
                       Patch(facecolor=C_FALTA, label="Faltante del día", alpha=0.55, hatch="//"),
                       plt.Line2D([0], [0], color=C_PLANEADO, ls="--", label="Planeado")],
              loc="upper left", fontsize=8.5, framealpha=0.9)
    plt.tight_layout()
    p = os.path.join(IMG, "d5_neteo.png")
    plt.savefig(p, dpi=150, bbox_inches="tight")
    plt.close()
    return p


# ─────────────────────────────────────────────────────────────────────────────
# Construir el .docx
# ─────────────────────────────────────────────────────────────────────────────
def build(paths):
    doc = Document()

    # estilos base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h(text, level=1):
        p = doc.add_heading(text, level=level)
        return p

    def para(text, bold=False, italic=False, size=11):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(size)
        return p

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        if bold_prefix:
            r = p.add_run(bold_prefix + " ")
            r.bold = True
        p.add_run(text)
        return p

    def img(path, width=6.5):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def caption(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    # Portada / titulo
    title = doc.add_heading("Cómo se descuenta el tiempo en el cálculo de asistencia", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("MundoVs — módulo RRHH · Rediseño 2026-07-27")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    para("Este documento explica, con gráficas, cómo se descuenta el tiempo desde los marcajes "
         "del reloj hasta el tiempo acreditado que se paga. Aplica a los dos modos diarios "
         "(Marcaje de Reloj y EntradaSalida) y al neteo semanal.")

    # 1. Conceptos
    h("1. Conceptos básicos", 1)
    bullet("tiempo entre la primera y la última marca del reloj.", "Bruto —")
    bullet("cada par intermedio (salida → regreso). Por defecto es pausa (se descuenta). Si el operador la edita como Trabajo, no se descuenta. Los descansos pagados no se descuentan. Si no hay par, es trabajo continuo (aunque el descanso esté planeado).", "Pausa —")
    bullet("Bruto menos las pausas no pagadas. Es lo que el reloj dice que se trabajó.", "Trabajado —")
    bullet("JornadaNeta programada (jornada bruta menos descansos no pagados), o el valor que el operador ponga en el modal (override).", "Planeado —")
    bullet("Max(0, Trabajado − Planeado). Lo que se trabajó encima de lo planeado.", "Extra detectado —")
    bullet("Max(0, Planeado − Trabajado − Permiso). Lo que faltó trabajar y no cubre un permiso.", "Faltante —")
    bullet("Min(Trabajado + Permiso, Planeado). Lo que se paga por el día (sin extra). El permiso cubre el faltante, no suma sobre la jornada.", "Tiempo Acreditado —")
    bullet("ausencia autorizada (real). Cubre el faltante. El tiempo perdonado (gracia que condona sin ser ausencia) NO aplica. La compensación de horas NO aplica.", "Permiso —")

    # 2. Flujo general
    h("2. Flujo general del descuento", 1)
    para("El descuento sigue siempre el mismo camino en los dos modos. La diferencia entre modos "
         "está en cómo se calcula el Trabajado y el Extra (ver sección 4).")
    img(paths["flujo"], width=6.8)
    caption("Figura 1. Flujo del descuento: Bruto → Trabajado (restando pausas) → comparado con Planeado → Extra/Faltante → Tiempo Acreditado.")

    # 3. Linea de tiempo
    h("3. Línea de tiempo: cómo se descuenta el descanso", 1)

    h("3.1 Trabajo continuo (sin par) — Marcaje de Reloj", 2)
    para("Caso del usuario: marcajes 11:10 (entrada) y 18:43 (salida), turno 11:30–19:00, con un "
         "descanso planeado D1 de 14:00 a 14:15 (15 min, no pagado). El descanso NO se marcó en el "
         "reloj y no hay par intermedio.")
    img(paths["t_marcaje"], width=6.8)
    para("Como no hay marca de pausa, el reloj indica trabajo continuo. El descanso planeado NO se "
         "descuenta (se trabajó). Resultado:")
    bullet("Bruto = 453 min (11:10 → 18:43)", "•")
    bullet("Pausa = 0 (sin par → trabajo continuo, aunque D1 esté planeado)", "•")
    bullet("Trabajado = 453 min", "•")
    bullet("Planeado = 435 min (jornada 450 − D1 no pagado 15)", "•")
    bullet("Extra detectado = 453 − 435 = 18 min", "•")
    bullet("Tiempo Acreditado = Min(453, 435) = 435 min", "•")

    h("3.2 Con pares intermedios (pausas reales) — caso Aralim", 2)
    para("Turno 8:00–18:15 con dos descansos planeados (D1 10:00–10:30, D2 14:00–14:45). El reloj "
         "registró seis marcas. D1 se marcó formalmente (11:03/11:30); las marcas 14:00/14:44 "
         "cayeron en la ventana del D2 pero sin clasificar — son un par intermedio, así que por "
         "defecto es pausa.")
    img(paths["t_aralim"], width=6.8)
    para("Aquí SÍ hay pares intermedios, así que se descuenta la pausa real (no la planeada):")
    bullet("Bruto = 669 min (7:09 → 18:18)", "•")
    bullet("Pausa D1 = 27 min (11:03 → 11:30, marcado)", "•")
    bullet("Pausa D2 = 44 min (14:00 → 14:44, par intermedio → pausa por defecto)", "•")
    bullet("Trabajado = 669 − 27 − 44 = 598 min", "•")
    bullet("Planeado = 540 min (jornada 615 − 30 − 45)", "•")
    bullet("Extra detectado = 598 − 540 = 58 min", "•")
    para("Regla clave: sin par → trabajo continuo (no se descuenta nada, aunque esté planeado); "
         "con par → se descuenta la pausa real, salvo que el operador edite el segmento como Trabajo.",
         italic=True)

    # 4. Barras acreditado/extra
    h("4. Trabajado vs Planeado → Acreditado + Extra", 1)
    para("Una vez que se tiene el Trabajado (reloj) y el Planeado, se comparan para sacar el Extra "
         "o el Faltante, y con el permiso se arma el Tiempo Acreditado.")
    img(paths["bars"], width=6.8)
    caption("Figura 3. Izquierda: trabajó más del planeado → Acreditado (hasta Planeado) + Extra. "
            "Derecha: trabajó menos; el permiso cubre parte del faltante y el resto queda como Faltante.")

    # 5. Diferencias entre modos
    h("5. Diferencia entre los dos modos", 1)
    tabla = doc.add_table(rows=1, cols=3)
    tabla.style = "Light Grid Accent 1"
    hdr = tabla.rows[0].cells
    hdr[0].text = "Aspecto"
    hdr[1].text = "Marcaje de Reloj"
    hdr[2].text = "EntradaSalida (default)"
    filas = [
        ("Usa horarios del turno", "No", "Sí"),
        ("Retardo / salida anticipada", "No se reportan", "Sí, informativo"),
        ("Umbral de extra (15 min)", "No aplica", "Sí (<15 → 0)"),
        ("Descanso NO marcado", "No se descuenta (trabajo continuo si no hay par)", "Descuenta el programado"),
        ("Par intermedio", "Pausa por defecto (descuenta real)", "Pausa por defecto (descuenta real)"),
        ("Bloque suelto previo/posterior", "Cuenta como extra", "Cuenta como extra + RequiereRevision"),
        ("Estatus", "Por Faltante/Extra", "Falta > Retardo > Salida anticipada > Normal"),
    ]
    for a, b, c in filas:
        row = tabla.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c
    for row in tabla.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.5)

    # 6. Neteo semanal
    h("6. Neteo semanal (liquidación)", 1)
    para("El extra se aprueba por semana, no por día. El neteo suma toda la semana y resta "
         "faltantes contra extras dentro de la misma semana.")
    img(paths["neteo"], width=6.8)
    caption("Figura 4. Neteo semanal: el extra del Lunes y Viernes se compensa con el faltante del "
            "Martes. El extra real que se aprueba es el neto de la semana (limitado al detectado).")
    bullet("Extra semanal (neto) = Max(0, ∑Trabajado − ∑Planeado).", "•")
    bullet("El operador aprueba hasta ese neto (tope = detectado).", "•")
    bullet("Factor: configurable + override del operador; afecta pago Y banco.", "•")
    bullet("Distribución: PagarTodo / MitadMitad / BancoTodo.", "•")
    bullet("Pago total = ∑ Tiempo Acreditado (días) + extra aprobado a pago × factor.", "•")

    # Guardar
    out_docx = os.path.join(OUT, "Calculo_Asistencia_Descuento_Tiempo.docx")
    doc.save(out_docx)
    return out_docx


if __name__ == "__main__":
    paths = {
        "t_marcaje": d1_timeline_marcaje(),
        "t_aralim": d2_timeline_aralim(),
        "bars": d3_bars(),
        "flujo": d4_flujo(),
        "neteo": d5_neteo(),
    }
    docx_path = build(paths)
    print("DOCX:", docx_path)
    for k, v in paths.items():
        print(k, v)